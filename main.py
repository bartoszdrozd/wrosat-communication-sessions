import os
import re
import pandas as pd
from access_functions import access_times, access_calculation
from docx import Document

document = Document()

pathName = os.getcwd()
csv_location = pathName + '/csvfiles'
fileNames = os.listdir(csv_location)
fileNames.sort()

csv_counter = 0
for fileName in fileNames:
	file_to_search = 'csvfiles/' + fileName
	df = pd.read_csv(file_to_search, index_col='Access')
	file_name = re.search("[A-Z]{3}\\d{0,3}", file_to_search).group(0)
	column = df['Duration (sec)']
	max_value = int(column.max())
	duration_list = list(range(50, max_value+50, 50))
	item_index = [0] * len(duration_list)
	if csv_counter < len(fileNames):
		access_times(df, item_index, duration_list, file_name)
		analysis_text = f"""
		Dane dla orbity: {file_name}
		Maksymalna długość sesji komunikacyjnej: {df['Duration (sec)'].max()}\n \
		Mediana czasów sesji komunikacyjnych: {df['Duration (sec)'].median()}\n \
		Średnia długość sesji komunikacyjnej: {df['Duration (sec)'].mean()}"""
		document.add_paragraph(analysis_text)
		document.add_picture(f'figures/{file_name}.png')
		access_calculation(df, file_name)
		document.add_picture(f'figures/length_of_silence_for_{file_name}.png')
		document.save('raport.docx')
		csv_counter += 1
		continue
